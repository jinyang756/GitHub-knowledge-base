#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import argparse
import zipfile
import shutil
from datetime import datetime

# 确保中文显示正常
import sys
import docx
from docx.shared import Inches

# 解析命令行参数
parser = argparse.ArgumentParser(description='将DOCX文件转换为Markdown格式')
parser.add_argument('input_file', help='输入的DOCX文件路径')
parser.add_argument('-o', '--output_dir', help='输出目录，默认与输入文件同目录')
parser.add_argument('--image_dir', help='图片保存目录，默认在输出目录下的images文件夹')
parser.add_argument('--verbose', action='store_true', help='显示详细转换过程')

args = parser.parse_args()

# 输入文件路径
input_file = os.path.abspath(args.input_file)

# 检查输入文件是否存在
if not os.path.exists(input_file):
    print(f"错误：文件 '{input_file}' 不存在")
    sys.exit(1)

# 检查文件扩展名
file_ext = os.path.splitext(input_file)[1].lower()
if file_ext != '.docx':
    print(f"错误：仅支持 .docx 格式文件，当前文件格式为 {file_ext}")
    sys.exit(1)

# 设置输出目录
if args.output_dir:
    output_dir = os.path.abspath(args.output_dir)
else:
    output_dir = os.path.dirname(input_file)

# 确保输出目录存在
os.makedirs(output_dir, exist_ok=True)

# 设置图片保存目录
if args.image_dir:
    image_dir = os.path.abspath(args.image_dir)
else:
    image_dir = os.path.join(output_dir, 'images')

# 确保图片目录存在
os.makedirs(image_dir, exist_ok=True)

# 输出文件路径
base_name = os.path.splitext(os.path.basename(input_file))[0]
output_file = os.path.join(output_dir, f'{base_name}.md')

# 记录开始转换时间
start_time = datetime.now()
print(f"开始转换 '{input_file}' 到 '{output_file}'")

# 加载文档
try:
    doc = docx.Document(input_file)
except Exception as e:
    print(f"加载文档时出错：{e}")
    sys.exit(1)

# 存储Markdown内容
markdown_content = []

# 存储脚注
footnotes = []

# 图片计数
image_count = 0

# 段落格式处理函数

def get_paragraph_style(paragraph):
    """获取段落样式"""
    if paragraph.style.name.startswith('Heading') or paragraph.style.name.startswith('标题'):
        # 提取标题级别
        level_match = re.search(r'\d+', paragraph.style.name)
        if level_match:
            level = int(level_match.group())
            return {'type': 'heading', 'level': min(level, 6)}  # 最多支持6级标题
    elif paragraph.style.name.startswith('List') or paragraph.style.name.startswith('列表'):
        return {'type': 'list', 'ordered': paragraph.style.name.find('编号') != -1}
    elif paragraph.style.name.startswith('Quote') or paragraph.style.name.startswith('引用'):
        return {'type': 'quote'}
    elif paragraph.style.name.startswith('Code') or paragraph.style.name.startswith('代码'):
        return {'type': 'code'}
    
    # 检查是否是无序列表（通过段落前的符号判断）
    if paragraph.text.startswith(('• ', '● ', '○ ', '□ ', '- ', '* ')):
        return {'type': 'list', 'ordered': False}
    
    # 检查是否是有序列表（通过段落前的数字判断）
    if re.match(r'^\d+\.?\s+', paragraph.text):
        return {'type': 'list', 'ordered': True}
    
    return {'type': 'normal'}

def extract_hyperlinks(paragraph):
    """提取段落中的超链接"""
    text = paragraph.text
    # 这里简化处理，实际docx库不直接支持超链接提取
    # 在实际应用中，可能需要使用更复杂的XML解析方式
    return text

def process_run(run):
    """处理文本运行（Run）"""
    text = run.text
    
    # 处理特殊字符
    text = text.replace('\u200b', '')  # 零宽空格
    text = text.replace('\u00a0', ' ')  # 非断行空格
    
    # 处理格式
    if run.bold:
        text = f'**{text}**'
    if run.italic:
        text = f'*{text}*'
    if run.underline:
        text = f'__{text}__'
    
    return text

def extract_images(paragraph, image_dir, base_name):
    """提取段落中的图片"""
    global image_count
    images = []
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            # 获取图片数据
            image_data = rel.target_part._blob
            # 确定图片格式
            content_type = rel.target_part.content_type
            if 'png' in content_type:
                ext = '.png'
            elif 'jpeg' in content_type:
                ext = '.jpg'
            elif 'gif' in content_type:
                ext = '.gif'
            else:
                ext = '.png'  # 默认使用png格式
            
            # 保存图片
            image_filename = f'{base_name}_image_{image_count}{ext}'
            image_path = os.path.join(image_dir, image_filename)
            
            try:
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                
                # 构建Markdown图片引用路径
                # 使用相对路径引用图片
                relative_image_path = os.path.relpath(image_path, output_dir)
                # 确保路径使用正斜杠
                relative_image_path = relative_image_path.replace('\\', '/')
                
                images.append(f'![图片 {image_count}]({relative_image_path})')
                
                if args.verbose:
                    print(f"  提取图片: {image_filename}")
            except Exception as e:
                if args.verbose:
                    print(f"  提取图片失败: {e}")
                images.append(f'[图片 {image_count} (提取失败)]')
    
    return images

def process_table(table):
    """处理表格"""
    markdown_table = []
    
    # 获取表格的行数和列数
    rows_count = len(table.rows)
    if rows_count == 0:
        return ''
    
    cols_count = len(table.rows[0].cells)
    
    # 处理表头
    header_cells = table.rows[0].cells
    header = '| ' + ' | '.join([cell.text.strip() for cell in header_cells]) + ' |'
    markdown_table.append(header)
    
    # 添加分隔线
    separator = '| ' + ' | '.join(['---'] * cols_count) + ' |'
    markdown_table.append(separator)
    
    # 处理表格内容
    for row in table.rows[1:]:
        cells = row.cells
        row_content = '| ' + ' | '.join([cell.text.strip() for cell in cells]) + ' |'
        markdown_table.append(row_content)
    
    return '\n'.join(markdown_table)

def extract_footnote(reference_id):
    """提取脚注内容"""
    # 这里简化处理，实际docx库不直接支持脚注提取
    # 在实际应用中，可能需要使用更复杂的XML解析方式
    return f"[脚注 {reference_id}]"

# 处理文档内容
previous_paragraph_type = 'normal'

for i, paragraph in enumerate(doc.paragraphs):
    # 跳过空段落
    if not paragraph.text.strip():
        if previous_paragraph_type != 'empty':  # 避免连续的空行
            markdown_content.append('')
            previous_paragraph_type = 'empty'
        continue
    
    # 获取段落样式
    style_info = get_paragraph_style(paragraph)
    
    # 处理段落文本
    paragraph_text = ''
    for run in paragraph.runs:
        paragraph_text += process_run(run)
    
    # 去除多余的空白字符
    paragraph_text = re.sub(r'\s+', ' ', paragraph_text).strip()
    
    # 根据样式处理段落
    if style_info['type'] == 'heading':
        # 标题处理
        level = style_info['level']
        markdown_content.append(f"{'#' * level} {paragraph_text}")
    elif style_info['type'] == 'list':
        # 列表处理
        if style_info['ordered']:
            # 有序列表，保留原始编号
            ordered_match = re.match(r'(\d+\.?\s+)(.*)', paragraph_text)
            if ordered_match:
                number = ordered_match.group(1)
                content = ordered_match.group(2)
                # 确保编号对齐，使用数字+点+空格格式
                normalized_number = re.search(r'\d+', number).group() + '. '
                markdown_content.append(f"{normalized_number}{content}")
            else:
                markdown_content.append(f"1. {paragraph_text}")
        else:
            # 无序列表
            # 移除段落前的符号
            content = re.sub(r'^[•●○□-*]\s+', '', paragraph_text)
            markdown_content.append(f"- {content}")
    elif style_info['type'] == 'quote':
        # 引用处理
        markdown_content.append(f"> {paragraph_text}")
    elif style_info['type'] == 'code':
        # 代码块处理
        # 检查前一行是否已经是代码块开始
        if markdown_content and not markdown_content[-1].startswith('```'):
            markdown_content.append('```')
        markdown_content.append(paragraph_text)
        # 标记下一行可能需要结束代码块
        needs_code_end = True
    else:
        # 普通段落处理
        # 检查是否需要结束代码块
        if 'needs_code_end' in locals() and needs_code_end:
            markdown_content.append('```')
            del needs_code_end
        
        # 添加普通段落
        markdown_content.append(paragraph_text)
    
    # 提取并处理图片
    images = extract_images(paragraph, image_dir, base_name)
    if images:
        markdown_content.extend(images)
    
    previous_paragraph_type = style_info['type']

# 处理表格
for table in doc.tables:
    table_content = process_table(table)
    if table_content:
        markdown_content.append('')
        markdown_content.append(table_content)
        markdown_content.append('')

# 处理脚注
if footnotes:
    markdown_content.append('')
    markdown_content.append('## 脚注')
    for i, footnote in enumerate(footnotes, 1):
        markdown_content.append(f'[{i}]: {footnote}')

# 最后检查是否需要结束代码块
if 'needs_code_end' in locals() and needs_code_end:
    markdown_content.append('```')

# 清理多余的分隔符和格式问题
cleaned_content = []
for line in markdown_content:
    # 移除多余的分隔符（超过5个连字符的行）
    if re.match(r'^-{5,}$', line):
        continue
    
    # 移除行尾的多余空格
    line = line.rstrip()
    
    # 添加清理后的行
    cleaned_content.append(line)

# 合并处理后的内容
final_content = '\n'.join(cleaned_content)

# 保存Markdown文件
try:
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(final_content)
    print(f"转换完成，已保存到 '{output_file}'")
    
    if image_count > 0:
        print(f"共提取 {image_count} 张图片到 '{image_dir}'")
        
    # 计算转换时间
    end_time = datetime.now()
    duration = (end_time - start_time).total_seconds()
    print(f"转换耗时: {duration:.2f} 秒")
except Exception as e:
    print(f"保存文件时出错：{e}")
    sys.exit(1)

# 输出转换统计信息
print(f"\n转换统计：")
print(f"- 原文档：{input_file}")
print(f"- 输出文件：{output_file}")
print(f"- 段落数：{len(doc.paragraphs)}")
print(f"- 表格数：{len(doc.tables)}")
print(f"- 提取图片数：{image_count}")

# 提示用户检查转换结果
print(f"\n请检查转换结果，可能需要手动调整一些格式问题。特别是对于复杂的表格、脚注和特殊格式。")